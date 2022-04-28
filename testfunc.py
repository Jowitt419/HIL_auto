def pplt():
    import pandas as pd
    from nptdms import TdmsFile as td
    import matplotlib.pyplot as plt
    import matplotlib as mpl
    mpl.rcParams['font.sans-serif'] = ['SimHei']
    mpl.rcParams['axes.unicode_minus'] = False
    # from matplotlib import gridspec
    # import time
    from docx import Document
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import re



    out_fp = r'C:/Users/jiahui/Desktop/HIL NI/HIL_auto/TestStand_log/ts_log.tdms'
    out_png = r'C:/Users/jiahui/Desktop/HIL NI/HIL_auto/HIL试验结果示例.png'
    docx_fp = r'C:/Users/jiahui/Desktop/HIL NI/HIL_auto/HIL试验报告示例.docx'
    df = pd.DataFrame()
    
    df_raw = pd.DataFrame()
    with td.open(out_fp) as tdms_file:
        print(tdms_file)
        metadata = td.read_metadata(out_fp)
        df_pro = pd.DataFrame([metadata.properties.values()], columns=metadata.properties.keys())
        df_raw = td(out_fp).as_dataframe()
        def postprocessing(df_raw, rate):
            col1 = df_raw.columns.to_list()
            # col2 = [n.replace('/\'VeriStand Channels\'/\'Aliases/', '').replace('\'', '') for n in col1]
            col2 = [re.sub(r'^.*Aliases/', '', n).replace('\'', '') for n in col1]#

            df_out = df_raw.rename(columns = dict(zip(col1, col2)))
            df_out['time'] = df_out.index*1/rate
            return df_out
        df = postprocessing(df_raw, rate = df_pro['Specified Rate [Hz]'][0])
    # print(df.head())
    # df.to_csv('temp.csv')
    
    fig = plt.figure(figsize = (6, 5), dpi = 200)
    gs = fig.add_gridspec(2,1, height_ratios=(2,1), hspace=0)
    ax1 = fig.add_subplot(gs[0])
    ax1.plot(df['time'], df['DesiredRPM'], color = 'green', label = 'DesiredRPM')
    ax1.plot(df['time'], df['ActualRPM'], color = 'blue', label = 'ActualRPM')
    ax1.set_ylabel('Engine RPM', fontsize = 12)
    plt.xticks(fontsize = 12)
    plt.yticks(fontsize = 12)
    plt.legend(loc = 'best', fontsize = 12)
    plt.grid()
    
    ax2 = fig.add_subplot(gs[1])
    ax2.plot(df['time'], df['EngineTemp'], color = 'orange', label = 'EngineTemp')
    ax2.set_ylabel('EngineTemp °C', fontsize = 12)
    ax2.set_xlabel('time (s)', fontsize = 12)
    plt.xticks(fontsize = 12)
    plt.yticks(fontsize = 12)
    plt.legend(loc = 'best', fontsize = 12)
    plt.grid()
    plt.suptitle('HIL试验结果示例', fontsize = 14)
    plt.savefig(fname = out_png)

    # # return True
        
    # def report_gen(df):#df = df_out, out_png = out_png, docx_fp = docx_fp
    #     document = Document()
    #     document.add_heading('HIL试验报告示例', level = 1)
    #     p = document.add_paragraph(
    #     'VCB基频', style='List Number')
    #     p.add_run('本次试车发动机最高温度为：' + str(round(float(df['EngineTemp'].max()),2)) + '°C@第' 
    #               + str(round(float(df.loc[df['EngineTemp'] == df['EngineTemp'].max(), 'time']),2)) + '秒')
    #     document.add_picture(out_png, width=Inches(6), height=Inches(5))#
        
    #     records = (
    #         ('1000', '999', '-1'),
    #         ('2000', '2001', '+1'),
    #         ('3000', '3002', '+2')
    #     )
        
    #     table = document.add_table(rows=1, cols=3)
    #     # table.allow_autofit = True
    #     table.autofit = False#True
    #     hdr_cells = table.rows[0].cells
    #     hdr_cells[0].text = '期望值'
    #     # hdr_cells[0].width = Inches(0.4)
    #     hdr_cells[1].text = '实际值'
    #     # hdr_cells[1].width = Inches(0.4)
    #     hdr_cells[2].text = '误差'
    #     # hdr_cells[2].width = Inches(0.4)
        
    #     for exp, rel, bias in records:
    #         row_cells = table.add_row().cells
    #         row_cells[0].text = str(exp)
    #         row_cells[1].text = str(rel)
    #         row_cells[2].text = str(bias)
        
    #     table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #     table.style = 'Table Grid'
    #     # document.add_page_break()
    #     document.save(docx_fp)
        
    # report_gen(df)
    a = 1+2
    return a
#
# jh = Report_Gentor()
# jh.plt_HIL()
# jh.report_gen()
