# import os
# from engine_demo.engine_demo_basic import run_engine_demo#examples.
# from niveristand import run_py_as_rtseq
from niveristand.errors import RunError
from niveristand.legacy import NIVeriStand
from NationalInstruments.VeriStand.Data import DoubleValue
from niveristand.library import wait
# import numpy as np
import pandas as pd
from nptdms import TdmsFile as td
import matplotlib.pyplot as plt
import matplotlib as mpl
mpl.rcParams['font.sans-serif'] = ['SimHei']
mpl.rcParams['axes.unicode_minus'] = False
# from matplotlib import gridspec
import time
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
time_now = time.strftime("%m%d_%H%M%S", time.localtime())
#%%
def HIL_test(df, in_fp, out_fp):
    """Combines the legacy API with Python real-time sequences to run a deterministic test."""
    # # Ensures NI VeriStand is running.
    # NIVeriStand.LaunchNIVeriStand()
    # NIVeriStand.WaitForNIVeriStandReady()
    # # Uses the ClientAPI interface to get a reference to Workspace2
    # workspace = NIVeriStand.Workspace2("localhost")
    # engine_demo_path = os.path.join(os.path.expanduser("~public"), 'Documents', 'National Instruments',
    #                                 'NI VeriStand 2019', 'Examples', 'Stimulus Profile', 'Engine Demo',
    #                                 'Engine Demo.nivssdf')
    # # Deploys the system definition.
    # workspace.ConnectToSystem(engine_demo_path, True, 120000)
    GATEWAY = "localhost"
    TARGET = "Controller" 
    workspace = NIVeriStand.Workspace2(GATEWAY)
    workspace.ReconnectToSystem(TARGET, True, None, 60000)
    
    Alias_dict = workspace.GetAliasList()
    # print(Alias_dict)
    
    log_info = NIVeriStand.CreateLogInfo()
    log_info.file_path = out_fp#'VeriStand.tdms'#
    log_info.description = "Engine demo logging"
    log_info.rate = 10
    channels_paths_to_log = [
        "Aliases/EnginePower",
        "Aliases/DesiredRPM",
        "Aliases/ActualRPM",
        "Aliases/EngineTemp"]
    channels_to_log = [NIVeriStand.CreateLogChannel(path) for path in channels_paths_to_log]
    NIVeriStand.SetLogInfoChannels(log_info, channels_to_log)


    try:
        # Uses Python real-time sequences to run a test.
        # run_py_as_rtseq(run_engine_demo)
        workspace.SetSingleChannelValue('Aliases/EnginePower', False)
        workspace.SetSingleChannelValue('Aliases/DesiredRPM', 0)
        workspace.SetSingleChannelValue('Aliases/EnginePower', True)
        # workspace.SetSingleChannelValue('Aliases/EnvTemp', '40')
        
        workspace.StartDataLogging(configuration_name = 'Logging', logInfo = log_info)
        for i in range(len(df)):
            # workspace.SetSingleChannelValue('Aliases/DesiredRPM', df['N1(rpm)'][i])
            workspace.SetMultipleChannelValues(['Aliases/DesiredRPM', 'Aliases/EnginePower'], # Aliases/EnvTemp
                                               [df['N1(rpm)'][i], True])
                                               # [df['N1(rpm)'][i], df['EnvTemp'][i]])
            wait(DoubleValue(df['period'][i]))
        workspace.StopDataLogging(configuration_name = 'Logging')
        
        
        
        pt_ActualRPM = workspace.GetSingleChannelValue('Aliases/ActualRPM')
        pt_array = workspace.GetMultipleChannelValues(['Aliases/ActualRPM', 'Aliases/EngineTemp'])
        pt_what = workspace.GetChannelVectorValues('Aliases/ActualRPM')
        out_list = [Alias_dict, pt_ActualRPM, pt_array, pt_what]
        
        
        
        workspace.SetSingleChannelValue('Aliases/EnginePower', False)
        workspace.SetSingleChannelValue('Aliases/DesiredRPM', 0)
        
        
        
        print("Test Success")
    except RunError as e:
        print("Test Failed: %d -  %s" % (int(e.error.error_code), e.error.message))
    finally:
        # You can now disconnect from the system, so the next test can run.
        workspace.DisconnectFromSystem('', True)

    return out_list



def postprocessing(df_raw, rate):
    col1 = df_raw.columns.to_list()
    col2 = [n.replace('/\'Aliases\'/\'', '').replace('\'', '') for n in col1]
    df_out = df_raw.rename(columns = dict(zip(col1, col2)))
    df_out['time'] = df_out.index*1/rate
    return df_out

def plot_HIL(df, out_png):
    fig = plt.figure(figsize = (6, 5), dpi = 200)
    gs = fig.add_gridspec(2,1, height_ratios=(2,1), hspace=0)
    ax1 = fig.add_subplot(gs[0])
    ax1.plot(df['time'], df['DesiredRPM'], color = 'green', label = 'DesiredRPM')
    ax1.plot(df['time'], df['ActualRPM'], color = 'blue', label = 'ActualRPM')
    ax1.set_ylabel('Engine RPM', fontsize = 12)
    plt.xticks(fontsize = 12)
    plt.yticks(fontsize = 12)
    plt.legend(loc = 'best', fontsize = 12)
    plt.grid()
    
    ax2 = fig.add_subplot(gs[1])
    ax2.plot(df['time'], df['EngineTemp'], color = 'orange', label = 'EngineTemp')
    ax2.set_ylabel('EngineTemp °C', fontsize = 12)
    ax2.set_xlabel('time (s)', fontsize = 12)
    plt.xticks(fontsize = 12)
    plt.yticks(fontsize = 12)
    plt.legend(loc = 'best', fontsize = 12)
    plt.grid()
    plt.suptitle('HIL试验结果示例', fontsize = 14)
    plt.savefig(out_png)
    
    
def create_docx(df, out_png, docx_fp):
    document = Document()
    document.add_heading('HIL试验报告示例', level = 1)
    
    # p = document.add_paragraph('振动部分')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    
    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')
    
    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    p = document.add_paragraph(
        'VCB基频', style='List Number'
    )
    p.add_run('本次试车发动机最高温度为：' + str(round(float(df['EngineTemp'].max()),2)) + '°C@第' 
              + str(round(float(df.loc[df['EngineTemp'] == df['EngineTemp'].max(), 'time']),2)) + '秒')
    document.add_picture(out_png, width=Inches(6), height=Inches(5))#
    
    records = (
        ('1000', '999', '-1'),
        ('2000', '2001', '+1'),
        ('3000', '3002', '+2')
    )
    
    table = document.add_table(rows=1, cols=3)
    # table.allow_autofit = True
    table.autofit = False#True
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '期望值'
    # hdr_cells[0].width = Inches(0.4)
    hdr_cells[1].text = '实际值'
    # hdr_cells[1].width = Inches(0.4)
    hdr_cells[2].text = '误差'
    # hdr_cells[2].width = Inches(0.4)
    
    for exp, rel, bias in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(exp)
        row_cells[1].text = str(rel)
        row_cells[2].text = str(bias)
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # document.add_page_break()
    document.save(docx_fp)
#%%
if __name__ == '__main__':
    in_fp = './HIL_in.csv'
    out_fp = 'C:/Users/jiahui/Desktop/HIL NI/HIL_auto/' + time_now + '.tdms'
    out_png = 'C:/Users/jiahui/Desktop/HIL NI/HIL_auto/HIL试验结果示例' + time_now + '.png'
    docx_fp = 'C:/Users/jiahui/Desktop/HIL NI/HIL_auto/' + time_now + 'HIL试验报告示例.docx'
    
    
    df = pd.read_csv(in_fp, encoding = 'utf-8', header = [0], skiprows = [1, 2])
    df['period'] = df['timestamp'].diff().shift(periods = -1).fillna(4)
    out_list = HIL_test(df = df, in_fp = in_fp, out_fp = out_fp)
    df_raw = pd.DataFrame()
    with td.open(out_fp) as tdms_file:
        print(tdms_file)
        metadata = td.read_metadata(out_fp)
        df_pro = pd.DataFrame([metadata.properties.values()], columns=metadata.properties.keys())
        df_raw = td(out_fp).as_dataframe()
        
    if len(df_raw):
        df_out = postprocessing(df_raw, rate = df_pro['Specified Rate [Hz]'][0])
        plot_HIL(df = df_out, out_png = out_png)
        create_docx(df = df_out, out_png = out_png, docx_fp = docx_fp)
#%%



    
    